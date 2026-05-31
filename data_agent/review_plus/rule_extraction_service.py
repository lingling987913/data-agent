"""
Review-Plus 规则抽取服务

从审查规则文档中提取结构化检查项：
  - xlsx: openpyxl 逐行解析，识别表头，映射到 CheckItem
  - docx: 段落分割，按编号+项目+要求模式提取
"""

import logging
import re
from pathlib import Path

from data_agent.review_plus.schemas import ReviewPlusCheckItem

logger = logging.getLogger(__name__)

_XLSX_HEADER_MAP = {
    "序号": "item_no",
    "编号": "item_no",
    "检查项目": "title",
    "审查项目": "title",
    "检查项": "title",
    "项目": "title",
    "检查要求": "requirement_text",
    "审查要求": "requirement_text",
    "检查内容": "requirement_text",
    "要求": "requirement_text",
    "验收准则": "acceptance_criteria",
    "检查对象": "applicable_scope",
    "适用范围": "applicable_scope",
    "适用对象": "applicable_scope",
    "备注": "notes",
    "严重度": "severity",
}

_DOCX_NUMBERED_PATTERN = re.compile(
    r"^(\d+(?:\.\d+)*)\s*[、.．)\]]\s*(.+)"
)

# 匹配 "- CHECK-XXX 标题" 或 "- XXX-001 标题" 格式的检查项
_CHECK_ITEM_PATTERN = re.compile(
    r"^[-*]\s+((?:CHECK|[A-Z]{2,6})-[A-Za-z0-9_-]+)\s+(.+)"
)

# 匹配 [Severity: xxx] 标记
_SEVERITY_PATTERN = re.compile(r"\[Severity:\s*(critical|major|minor|info)\]", re.IGNORECASE)

_NON_ITEM_TITLE_RE = re.compile(
    r"^(密级|公开|产品保证工作检查单|文档名称|文件编号|编制|审核|批准|日期|目录|前言|封面)[:：]?"
)


def _looks_like_noise(title: str, requirement: str = "") -> bool:
    text = f"{title} {requirement}".strip()
    if not text:
        return True
    if _NON_ITEM_TITLE_RE.match(text):
        return True
    if len(text) < 6 and not re.search(r"检查|要求|符合|一致|分析|验证", text):
        return True
    return False


def extract_check_items_from_xlsx(
    file_path: str,
    material_name: str = "",
) -> list[ReviewPlusCheckItem]:
    """从 xlsx 文件中提取检查项。"""
    try:
        import openpyxl
    except ImportError:
        logger.warning("[RuleExtraction] openpyxl 未安装，无法解析 xlsx")
        return []

    items: list[ReviewPlusCheckItem] = []
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    except Exception as e:
        logger.warning(f"[RuleExtraction] 无法打开 xlsx: {e}")
        return []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_iter = ws.iter_rows(values_only=True)

        # 找表头行
        header_row = None
        header_map: dict[int, str] = {}
        for row in rows_iter:
            cells = [str(c).strip() if c else "" for c in row]
            matched = sum(1 for c in cells if c in _XLSX_HEADER_MAP)
            if matched >= 2:
                header_row = cells
                for idx, cell in enumerate(cells):
                    field = _XLSX_HEADER_MAP.get(cell)
                    if field:
                        header_map[idx] = field
                break

        if not header_map:
            continue

        # 逐行提取
        for row_num, row in enumerate(rows_iter, start=2):
            if row is None:
                continue
            cells = [str(c).strip() if c else "" for c in row]
            # 跳过空行
            if not any(cells):
                continue

            values: dict[str, str] = {}
            for idx, field in header_map.items():
                if idx < len(cells):
                    values[field] = cells[idx]

            # 至少要有 title 或 requirement_text
            if not values.get("title") and not values.get("requirement_text"):
                continue

            items.append(ReviewPlusCheckItem(
                item_no=values.get("item_no", ""),
                title=values.get("title", ""),
                requirement_text=values.get("requirement_text", ""),
                acceptance_criteria=values.get("acceptance_criteria", ""),
                applicable_scope=values.get("applicable_scope", ""),
                severity=values.get("severity", "minor"),
                source_material_name=material_name,
                source_sheet=sheet_name,
                source_row=row_num,
                source_quote=values.get("requirement_text", ""),
                confidence=0.9,
            ))

    wb.close()
    logger.info(f"[RuleExtraction] xlsx 提取完成: {len(items)} 条检查项, file={material_name}")
    return items


def extract_check_items_from_docx(
    file_path: str,
    material_name: str = "",
) -> list[ReviewPlusCheckItem]:
    """从 docx 文件中提取检查项（按编号段落模式）。"""
    try:
        import docx
    except ImportError:
        logger.warning("[RuleExtraction] python-docx 未安装，无法解析 docx")
        return []

    items: list[ReviewPlusCheckItem] = []
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        logger.warning(f"[RuleExtraction] 无法打开 docx: {e}")
        return []

    table_items = _extract_check_items_from_docx_tables(doc, material_name)
    if table_items:
        logger.info(f"[RuleExtraction] docx 表格提取完成: {len(table_items)} 条检查项, file={material_name}")
        return table_items

    current_item_no = ""
    current_title = ""
    current_requirement = ""
    row_num = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        row_num += 1
        m = _DOCX_NUMBERED_PATTERN.match(text)
        if m:
            # 保存上一条
            if current_title or current_requirement:
                if not _looks_like_noise(current_title, current_requirement):
                    items.append(ReviewPlusCheckItem(
                        item_no=current_item_no,
                        title=current_title,
                        requirement_text=current_requirement,
                        source_material_name=material_name,
                        source_row=row_num - 1,
                        source_quote=current_requirement,
                        confidence=0.7,
                    ))

            current_item_no = m.group(1)
            rest = m.group(2).strip()
            # 尝试分割标题和要求（以 "：" 或 ":" 分隔）
            if "：" in rest:
                parts = rest.split("：", 1)
            elif ":" in rest:
                parts = rest.split(":", 1)
            else:
                parts = [rest, ""]

            current_title = parts[0].strip()
            current_requirement = parts[1].strip() if len(parts) > 1 else ""
        else:
            # 追加到当前 requirement
            current_requirement = (current_requirement + " " + text).strip()

    # 保存最后一条
    if current_title or current_requirement:
        if not _looks_like_noise(current_title, current_requirement):
            items.append(ReviewPlusCheckItem(
                item_no=current_item_no,
                title=current_title,
                requirement_text=current_requirement,
                source_material_name=material_name,
                source_row=row_num,
                source_quote=current_requirement,
                confidence=0.7,
            ))

    logger.info(f"[RuleExtraction] docx 提取完成: {len(items)} 条检查项, file={material_name}")
    return items


def _extract_check_items_from_docx_tables(doc, material_name: str = "") -> list[ReviewPlusCheckItem]:
    items: list[ReviewPlusCheckItem] = []
    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])
        if not rows:
            continue

        header_idx = -1
        header_map: dict[int, str] = {}
        for idx, row in enumerate(rows[:5]):
            header_map = {}
            for col_idx, cell in enumerate(row):
                field = _XLSX_HEADER_MAP.get(cell.strip())
                if field:
                    header_map[col_idx] = field
            if len(header_map) >= 2:
                header_idx = idx
                break
        if header_idx < 0:
            continue

        for row_offset, row in enumerate(rows[header_idx + 1:], start=header_idx + 2):
            values: dict[str, str] = {}
            for col_idx, field in header_map.items():
                if col_idx < len(row):
                    values[field] = row[col_idx]
            title = values.get("title", "").strip()
            requirement = values.get("requirement_text", "").strip()
            if _looks_like_noise(title, requirement):
                continue
            items.append(ReviewPlusCheckItem(
                item_no=values.get("item_no", "").strip(),
                title=title,
                requirement_text=requirement,
                acceptance_criteria=values.get("acceptance_criteria", "").strip(),
                applicable_scope=values.get("applicable_scope", "").strip(),
                severity=values.get("severity", "minor").strip() or "minor",
                source_material_name=material_name,
                source_sheet=f"table-{table_index}",
                source_row=row_offset,
                source_quote=requirement or title,
                confidence=0.82,
            ))
    return items


def extract_check_items_from_text(
    content: str,
    material_name: str = "",
) -> list[ReviewPlusCheckItem]:
    """从纯文本内容中按编号模式或 CHECK-XXX 模式提取检查项。"""
    items: list[ReviewPlusCheckItem] = []
    if not content:
        return items

    current_item_no = ""
    current_title = ""
    current_requirement = ""
    row_num = 0

    # 跳过 markdown 标题行和元数据行（不以列表符号或编号开头）
    _SKIP_LINE = re.compile(r"^\s*(#|文档编号|基线版本|>|---|\*\*|```)")
    _EMPTY_ITEM = ("", "", "")

    for line in content.split("\n"):
        text = line.strip()
        if not text:
            # 空行触发保存当前累积项（仅在有编号时才保存，避免标题等被误收）
            if current_item_no and (current_title or current_requirement):
                items.append(_build_text_check_item(
                    current_item_no, current_title, current_requirement,
                    material_name, row_num,
                ))
                current_item_no = ""
                current_title = ""
                current_requirement = ""
            continue

        # 跳过标题、元数据等非检查项行
        if _SKIP_LINE.match(text):
            continue

        row_num += 1

        # 优先尝试 CHECK-XXX / REQ-XXX / DES-XXX 格式
        check_match = _CHECK_ITEM_PATTERN.match(text)
        if check_match:
            # 保存前一条
            if current_title or current_requirement:
                items.append(_build_text_check_item(
                    current_item_no, current_title, current_requirement,
                    material_name, row_num - 1,
                ))
            current_item_no = check_match.group(1)
            rest = check_match.group(2).strip()
            severity, cleaned = _extract_severity(rest)
            # 按中文冒号分割标题和要求
            if "：" in cleaned:
                parts = cleaned.split("：", 1)
            elif ":" in cleaned:
                parts = cleaned.split(":", 1)
            else:
                parts = [cleaned, ""]
            current_title = parts[0].strip()
            current_requirement = parts[1].strip() if len(parts) > 1 else ""
            if severity:
                current_requirement = _strip_severity_tag(current_requirement or rest)
            continue

        # 退化到原有编号模式
        m = _DOCX_NUMBERED_PATTERN.match(text)
        if m:
            if current_title or current_requirement:
                items.append(_build_text_check_item(
                    current_item_no, current_title, current_requirement,
                    material_name, row_num - 1,
                ))
            current_item_no = m.group(1)
            rest = m.group(2).strip()
            if "：" in rest:
                parts = rest.split("：", 1)
            elif ":" in rest:
                parts = rest.split(":", 1)
            else:
                parts = [rest, ""]
            current_title = parts[0].strip()
            current_requirement = parts[1].strip() if len(parts) > 1 else ""
        else:
            current_requirement = (current_requirement + " " + text).strip()

    if current_title or current_requirement:
        items.append(_build_text_check_item(
            current_item_no, current_title, current_requirement,
            material_name, row_num,
        ))

    return items


def _extract_severity(text: str) -> tuple[str, str]:
    """从文本中提取 [Severity: xxx] 标记，返回 (severity, cleaned_text)。"""
    m = _SEVERITY_PATTERN.search(text)
    if m:
        return m.group(1).lower(), _SEVERITY_PATTERN.sub("", text).strip()
    return "", text


def _strip_severity_tag(text: str) -> str:
    """移除 [Severity: xxx] 标记。"""
    return _SEVERITY_PATTERN.sub("", text).strip()


def _build_text_check_item(
    item_no: str,
    title: str,
    requirement_text: str,
    material_name: str,
    row_num: int,
    *,
    confidence: float = 0.6,
) -> ReviewPlusCheckItem:
    """构建文本提取的检查项，自动从 title/requirement 中提取 severity。"""
    combined = f"{title} {requirement_text}"
    severity_str, _ = _extract_severity(combined)
    clean_title = _strip_severity_tag(title)
    clean_requirement = _strip_severity_tag(requirement_text)
    if _looks_like_noise(clean_title, clean_requirement):
        return ReviewPlusCheckItem(
            item_no="",
            title="",
            requirement_text="",
            source_material_name=material_name,
            source_row=row_num,
            confidence=0.0,
        )
    return ReviewPlusCheckItem(
        item_no=item_no,
        title=clean_title,
        requirement_text=clean_requirement,
        source_material_name=material_name,
        source_row=row_num,
        source_quote=clean_requirement,
        severity=severity_str or "minor",
        confidence=confidence,
    )


def extract_check_items(
    file_path: str,
    material_name: str = "",
    content: str = "",
) -> list[ReviewPlusCheckItem]:
    """根据文件类型自动选择抽取方法。"""
    ext = Path(file_path).suffix.lower() if file_path else ""
    name_ext = Path(material_name).suffix.lower() if material_name else ext

    if name_ext in (".xlsx", ".xls") and file_path and Path(file_path).exists():
        return extract_check_items_from_xlsx(file_path, material_name)

    if name_ext in (".docx",) and file_path and Path(file_path).exists():
        return [
            item for item in extract_check_items_from_docx(file_path, material_name)
            if item.title or item.requirement_text
        ]

    # 其他格式：从 content 文本提取
    return [
        item for item in extract_check_items_from_text(content, material_name)
        if item.title or item.requirement_text
    ]
