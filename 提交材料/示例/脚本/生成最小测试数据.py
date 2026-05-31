#!/usr/bin/env python3
"""生成示例 01–04 的最小脱敏 测试数据（ywdata 不可用时的评审 fallback）。"""

from __future__ import annotations

from pathlib import Path

ROOT = Path(__file__).resolve().parents[3]
EXAMPLES = ROOT / "提交材料" / "示例"


def _write_qa_checklist(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("产品保证工作检查单（脱敏样例）", level=1)
    doc.add_paragraph("检查项 PA-001：飞轮可靠性分析是否覆盖关键失效模式。")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "编号"
    table.rows[0].cells[1].text = "检查内容"
    table.rows[1].cells[0].text = "PA-001"
    table.rows[1].cells[1].text = "设计报告是否引用任务书指标 REQ-FW-001。"
    doc.save(path)


def _write_task_book(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("月兔一号飞行器飞轮研制任务书（脱敏样例）", level=1)
    doc.add_paragraph("REQ-FW-001：飞轮角动量容量不小于 15 N·m·s。")
    doc.add_paragraph("REQ-FW-002：启动时间不大于 30 s。")
    doc.save(path)


def _write_design_report(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("飞轮可靠性安全性设计与分析报告（脱敏样例）", level=1)
    doc.add_paragraph("DES-FW-001 响应 REQ-FW-001：角动量设计值 16 N·m·s，满足任务书指标。")
    doc.add_paragraph("DES-FW-002 响应 REQ-FW-002：启动时间设计值 28 s。")
    doc.save(path)


def _write_xlsx(path: Path) -> None:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "文档检查需求"
    ws.append(["编号", "检查内容"])
    ws.append(["DOC-001", "任务书与设计报告编号体系一致。"])
    ws.append(["DOC-002", "验收报告与任务书指标可追溯。"])
    wb.save(path)
    wb.close()


def _write_pdf(path: Path) -> None:
    # 最小可读 PDF（纯文本流，供 pdftotext / 降级链冒烟）
    content = b"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj
4 0 obj<</Length 68>>stream
BT /F1 12 Tf 72 720 Td (CMG50 acceptance report - sanitized sample) Tj ET
endstream
endobj
5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000261 00000 n 
0000000380 00000 n 
trailer<</Size 6/Root 1 0 R>>
startxref
456
%%EOF
"""
    path.write_bytes(content)


def _ensure_dir(example: str) -> Path:
    d = EXAMPLES / example / "测试数据"
    d.mkdir(parents=True, exist_ok=True)
    return d


def main() -> None:
    files = {
        "月兔一号_产品保证检查单.docx": _write_qa_checklist,
        "月兔一号_飞轮研制任务书.docx": _write_task_book,
        "月兔一号_飞轮设计分析报告.docx": _write_design_report,
        "月兔一号_文档检查需求.xlsx": _write_xlsx,
    }

    for ex in ("01-多格式结构化", "04-规划与编排"):
        d = _ensure_dir(ex)
        for name, writer in files.items():
            writer(d / name)
            print(f"  {d / name}")

    d03 = _ensure_dir("03-跨文档指代")
    _write_task_book(d03 / "月兔一号_飞轮研制任务书.docx")
    _write_design_report(d03 / "月兔一号_飞轮设计分析报告.docx")
    print(f"  {d03}")

    d02 = _ensure_dir("02-PDF验收解析")
    pdf = d02 / "CMG50_验收报告.pdf"
    _write_pdf(pdf)
    print(f"  {pdf}")

    print("Done. 脱敏 测试数据 已写入各 示例/*/测试数据/。")


if __name__ == "__main__":
    main()
